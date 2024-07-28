import os  # 用于处理文件和目录路径
import pypandoc  # 用于文件格式转换
from docx import Document  # 用于处理Word文档
import shutil  # 用于删除文件和文件夹
import logging  # 用于记录日志

# 配置日志记录，设置日志格式和日志级别
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# 删除指定文件夹中不在排除列表中的文件和文件夹
def delete_existing_files_and_folders(folder, exclude_files):
    """
    删除指定文件夹中不在排除列表中的文件和文件夹。

    参数:
    folder (str): 要删除文件和文件夹的文件夹路径。
    exclude_files (list): 要保留的文件名列表。
    """
    try:
        for item_name in os.listdir(folder):  # 列出文件夹中的所有文件和文件夹
            item_path = os.path.join(folder, item_name)
            if item_name not in exclude_files:  # 检查是否在排除列表中
                if os.path.isfile(item_path):  # 检查是否为文件
                    os.remove(item_path)  # 删除文件
                    logging.info(f"已删除文件: {item_path}")
                elif os.path.isdir(item_path):  # 检查是否为文件夹
                    shutil.rmtree(item_path)  # 删除文件夹及其所有内容
                    logging.info(f"已删除文件夹: {item_path}")
    except Exception as e:
        logging.error(f"删除文件和文件夹时出错: {e}")

# 将文件转换为Markdown格式
def convert_file_md(source_file, output_file):
    """
    将文件转换为Markdown格式。

    参数:
    source_file (str): 源文件路径。
    output_file (str): 输出Markdown文件路径。
    """
    output_format = 'markdown'
    extract_media_path = os.path.splitext(output_file)[0] + '_media'  # 设置媒体文件的路径
    extra_args = ['--extract-media=' + extract_media_path, '--toc', '--standalone']
    output = pypandoc.convert_file(source_file, to=output_format, outputfile=output_file, extra_args=extra_args)
    replace_backslashes_in_md(output_file)
    logging.info("MD文件格式转换成功")

# 替换Markdown文件中的反斜杠为正斜杠
def replace_backslashes_in_md(md_file_path):
    """
    替换Markdown文件中的反斜杠为正斜杠。

    参数:
    md_file_path (str): Markdown文件路径。
    """
    with open(md_file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    updated_content = content.replace('\\', '/')
    with open(md_file_path, 'w', encoding='utf-8') as file:
        file.write(updated_content)

# 为Markdown文件添加缩进
def add_indents_to_md(md_file_path, docx_file_path):
    """
    为Markdown文件添加缩进。

    参数:
    md_file_path (str): Markdown文件路径。
    docx_file_path (str): 包含缩进信息的Word文档路径。
    """
    with open(md_file_path, 'r', encoding='utf-8') as file:
        md_content = file.readlines()

    docx = Document(docx_file_path)
    indents = []
    for para in docx.paragraphs:  # 获取每个段落的缩进信息
        left_indent = para.paragraph_format.left_indent
        first_line_indent = para.paragraph_format.first_line_indent
        hanging_indent = -first_line_indent if first_line_indent and first_line_indent < 0 else 0
        indents.append((left_indent, first_line_indent, hanging_indent))

    new_md_content = []
    indent_idx = 0
    for line in md_content:
        if line.strip() != "" and not line.startswith("[^"):  # 忽略注解行
            if indent_idx < len(indents):
                left_indent, first_line_indent, hanging_indent = indents[indent_idx]
                indent_text = ''
                if left_indent:
                    indent_text += '[left_indent]'
                if first_line_indent:
                    indent_text += '[first_line_indent]'
                if hanging_indent:
                    indent_text += '[hanging_indent]'
                new_md_content.append(indent_text + line)
                indent_idx += 1
            else:
                logging.warning(f"Markdown中的段落比DOCX中的多。跳过行的缩进: {line.strip()}")
                new_md_content.append(line)
        else:
            new_md_content.append(line)

    with open(md_file_path, 'w', encoding='utf-8') as file:
        file.writelines(new_md_content)

# 将Markdown文件转换为Word文档
def convert_md_to_docx(input_md_path, output_docx_path):
    """
    将Markdown文件转换为Word文档。

    参数:
    input_md_path (str): 输入Markdown文件路径。
    output_docx_path (str): 输出Word文档路径。
    """
    input_md_path = os.path.normpath(input_md_path)
    output_docx_path = os.path.normpath(output_docx_path)
    output = pypandoc.convert_file(input_md_path, 'docx', outputfile=output_docx_path)
    logging.info("Word文档格式转换成功")
    return output

# 新增功能: 替换缩进占位符为制表符
def replace_indent_placeholders(docx_file_path):
    """
    替换Word文档中的缩进占位符为制表符。

    参数:
    docx_file_path (str): Word文档路径。
    """
    # 打开要处理的DOCX文件
    doc = Document(docx_file_path)
    for para in doc.paragraphs:
        text = para.text

        # 用制表符替换左缩进占位符
        text = text.replace('[left_indent]', '\t')
        
        # 用制表符替换首行缩进占位符
        text = text.replace('[first_line_indent]', '\t')
        
        # 用制表符替换悬挂缩进占位符
        text = text.replace('[hanging_indent]', '\t')

        # 更新段落文本
        if text != para.text:
            para.clear()  # 清空段落的运行
            para.add_run(text)  # 添加新的运行

    # 保存修改后的DOCX文件
    doc.save(docx_file_path)
    logging.info(f"Indent placeholders replaced in {docx_file_path}")

# 主程序入口
if __name__ == '__main__':
    test_folder = 'Doc_conver_test'  # 测试文件夹
    test_files = os.listdir(test_folder)

    for file_name in test_files:
        source_file = os.path.join(test_folder, file_name)
        if file_name.endswith('.docx') and not file_name.endswith('_converted.docx'):  # 仅处理未转换的DOCX文件
            exclude_files = [file_name]  # 保留源文件
            delete_existing_files_and_folders(test_folder, exclude_files)  # 删除其他文件和文件夹

            output_md_file = os.path.splitext(source_file)[0] + '.md'
            output_docx_file = os.path.splitext(source_file)[0] + '_converted.docx'
            
            # 转换DOCX文件为MD文件
            convert_file_md(source_file, output_md_file)
            # 为MD文件添加缩进占位符
            add_indents_to_md(output_md_file, source_file)
            # 将MD文件转换为DOCX文件
            convert_md_to_docx(output_md_file, output_docx_file)
            # 替换缩进占位符
            replace_indent_placeholders(output_docx_file)  # 新增功能: 转换后替换缩进占位符
