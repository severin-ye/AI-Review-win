from docx import Document

def replace_indent_placeholders(docx_file_path):
    # 打开要处理的DOCX文件
    doc = Document(docx_file_path)
    for para in doc.paragraphs:
        text = para.text

        # 用制表符替换左缩进占位符
        while '[left_indent:' in text:
            start_idx = text.find('[left_indent:')
            end_idx = text.find('pt]', start_idx) + 3
            text = text[:start_idx] + '\t' + text[end_idx:]
        
        # 用制表符替换首行缩进占位符
        while '[first_line_indent:' in text:
            start_idx = text.find('[first_line_indent:')
            end_idx = text.find('pt]', start_idx) + 3
            text = text[:start_idx] + '\t' + text[end_idx:]
        
        # 用制表符替换悬挂缩进占位符
        while '[hanging_indent:' in text:
            start_idx = text.find('[hanging_indent:')
            end_idx = text.find('pt]', start_idx) + 3
            text = text[:start_idx] + '\t' + text[end_idx:]

        # 更新段落文本
        if text != para.text:
            para.clear()  # 清空段落的运行
            para.add_run(text)  # 添加新的运行


    # 保存修改后的DOCX文件
    doc.save(docx_file_path)
    print(f"Docx文件缩进占位符替换成功: {docx_file_path}")

# 创建测试DOCX文件
def create_test_docx(file_path):
    doc = Document()

    # first_line_indent测试
    doc.add_paragraph("[first_line_indent: 32.0pt] This paragraph has first line indent.")
    # hanging_indent测试
    doc.add_paragraph("[hanging_indent: 10.0pt] This paragraph has hanging indent.")

    doc.save(file_path)

# 运行测试
test_docx_path = 'test_indent_placeholders.docx'
create_test_docx(test_docx_path)
replace_indent_placeholders(test_docx_path)

# 检查结果
doc = Document(test_docx_path)
for para in doc.paragraphs:
    print(f"Paragraph text: {para.text}")
