"""M3 检索测试：知识库索引（LanceDB + BM25）+ 查询重写 + 3+3 混合检索。

- embedding 一律走 stub provider（settings embedding.provider=stub，确定性假向量，不加载模型）。
- LLM 一律 mock（monkeypatch app.rag.retrieve.chat_json），不发起真实请求。
- 环境隔离同 M2：导入 app 前设置 AI_REVIEW_DATA_DIR / AI_REVIEW_SEGMENTER=rule；
  与其他测试模块同跑时引擎绑定首个导入模块的数据目录，路径断言一律以 get_settings() 为准。
"""
import os
import tempfile
import time
from pathlib import Path

_tmp = tempfile.mkdtemp(prefix="ai-review-test-rag-")
os.environ.setdefault("AI_REVIEW_DATA_DIR", _tmp)
os.environ.setdefault("AI_REVIEW_SEGMENTER", "rule")

import pytest  # noqa: E402
from docx import Document as DocxDocument  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

from app.core.config import get_settings  # noqa: E402
from app.llm.client import LLMNotConfiguredError  # noqa: E402
from app.main import app  # noqa: E402
from app.rag import store  # noqa: E402
from app.rag.index import chunk_id, load_kb_file, split_chunks  # noqa: E402
from app.rag.retrieve import rrf_fuse  # noqa: E402

DATA_DIR = get_settings().data_dir

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TXT_MIME = "text/plain"

KB_TEXT = """中国高血压防治指南：高血压定义为在未使用降压药物的情况下，非同日三次测量诊室血压，收缩压≥140mmHg和（或）舒张压≥90mmHg。根据血压升高水平，将高血压分为1级、2级和3级。降压治疗的根本目标是降低心、脑、肾与血管并发症和死亡的总危险。

2型糖尿病防治要点：二甲双胍是2型糖尿病的一线首选用药，无禁忌证者应一直保留在治疗方案中。常用起始剂量500mg每日两次，随餐服用可减少胃肠道反应。糖化血红蛋白控制目标一般为低于7.0%，需根据患者年龄与并发症个体化调整。

阿司匹林肠溶片说明：阿司匹林通过不可逆抑制环氧化酶，阻断血栓素A2生成，从而抑制血小板聚集，广泛用于动脉粥样硬化性心血管疾病的二级预防。常用剂量为每日75至100mg，长期服用需注意消化道出血风险，禁用于活动性溃疡患者。

慢性阻塞性肺疾病诊治：慢阻肺以持续气流受限为特征，肺功能检查吸入支气管扩张剂后FEV1/FVC小于0.70可确诊。稳定期首选长效支气管扩张剂，急性加重期可短程使用全身糖皮质激素并给予控制性氧疗，目标血氧饱和度88%至92%。

胰岛素使用规范：1型糖尿病患者需终身胰岛素替代治疗。基础胰岛素起始剂量通常为每日每公斤体重0.1至0.2单位，根据空腹血糖每3天调整一次。低血糖是最常见的不良反应，血糖低于3.9mmol/L需立即补充15克快速碳水化合物。

他汀类药物应用：阿托伐他汀通过抑制HMG-CoA还原酶降低低密度脂蛋白胆固醇，是动脉粥样硬化性心血管疾病一级与二级预防的基石用药。常规起始剂量10至20mg每晚一次，用药期间需监测肝酶与肌酸激酶，警惕横纹肌溶解。
"""


def _build_medical_docx(path: Path) -> None:
    """医学测试文档：正文句子 + 2 个表格（首个按审查意见表移除，第二个 → 占位符）+ 参考文献区。"""
    doc = DocxDocument()
    doc.add_heading("一、病例资料", level=1)
    doc.add_paragraph("患者男性，六十二岁，确诊高血压三年，规律服用氨氯地平控制血压。")
    doc.add_paragraph("既往有2型糖尿病病史五年，目前口服二甲双胍治疗，血糖控制尚可。")
    table1 = doc.add_table(rows=2, cols=2)  # 首个表格：审查意见表（has_review_table=Y 时移除）
    table1.cell(0, 0).text = "审查意见"
    table2 = doc.add_table(rows=2, cols=2)  # 第二个表格 → 占位符 [{表格不予审校_1}]
    table2.cell(0, 0).text = "检查项目"
    doc.add_paragraph("入院查体血压150/95mmHg，心肺听诊未见明显异常。")
    doc.add_heading("参考文献", level=2)
    doc.add_paragraph("[1]: 中国高血压防治指南修订委员会. 中华心血管病杂志, 2019.")
    doc.save(str(path))


@pytest.fixture(scope="module")
def client():
    with TestClient(app) as c:
        # stub embedding + 开启检索 + 固定首个表格为审查意见表（settings 实时读库，立即生效）
        assert c.put(
            "/api/settings",
            json={
                "embedding.provider": "stub",
                "retrieve.enabled": True,
                "docx.has_review_table": "Y",
            },
        ).status_code == 200
        yield c


@pytest.fixture(scope="module")
def docx_path() -> Path:
    path = Path(_tmp) / "medical.docx"
    _build_medical_docx(path)
    return path


def _wait_kb_status(client: TestClient, kb_id: str, timeout: float = 90) -> dict:
    deadline = time.time() + timeout
    while time.time() < deadline:
        docs = client.get("/api/kb/documents").json()
        row = next((d for d in docs if d["id"] == kb_id), None)
        if row and row["status"] in ("indexed", "failed"):
            return row
        time.sleep(0.5)
    raise TimeoutError(f"知识库索引超时: {kb_id}")


def _upload_kb(client: TestClient, name: str, content: bytes) -> dict:
    resp = client.post("/api/kb/documents", files={"file": (name, content, TXT_MIME)})
    assert resp.status_code == 201, resp.text
    row = _wait_kb_status(client, resp.json()["id"])
    assert row["status"] == "indexed", row
    return row


# ---------- 1. 切块 / 稳定 ID / 加载器 ----------


def test_chunking_stable_ids_and_loaders() -> None:
    chunks = split_chunks(KB_TEXT)
    assert len(chunks) >= 2  # 六段医学文本按 500 字切块
    assert all(len(c) <= 500 for c in chunks)
    assert all(c.strip() for c in chunks)

    doc_hash = "ab" * 32
    assert chunk_id(doc_hash, 0) == chunk_id(doc_hash, 0)  # 稳定
    assert chunk_id(doc_hash, 0) != chunk_id(doc_hash, 1)
    assert len(chunk_id(doc_hash, 0)) == 16

    # txt / csv / docx 加载器（pdf 走 pypdf，单测不造二进制，见 M3 文档已知问题）
    txt_path = Path(_tmp) / "a.txt"
    txt_path.write_text("第一行。\n第二行。", encoding="utf-8")
    assert "第二行" in load_kb_file(txt_path)

    csv_path = Path(_tmp) / "b.csv"
    csv_path.write_text("药品,剂量\n二甲双胍,500mg", encoding="utf-8")
    csv_text = load_kb_file(csv_path)
    assert "二甲双胍" in csv_text and "500mg" in csv_text

    docx_kb = Path(_tmp) / "c.docx"
    d = DocxDocument()
    d.add_paragraph("知识库段落文本。")
    d.save(str(docx_kb))
    assert "知识库段落文本" in load_kb_file(docx_kb)


# ---------- 2. RRF 融合数学正确性 ----------


def test_rrf_fuse_math() -> None:
    rankings = [["a", "b", "c"], ["b", "a", "d"]]
    scores = rrf_fuse(rankings, k=60)
    # score(d) = Σ 1/(60 + rank)，rank 从 1 起
    assert scores["a"] == pytest.approx(1 / 61 + 1 / 62)
    assert scores["b"] == pytest.approx(1 / 62 + 1 / 61)
    assert scores["c"] == pytest.approx(1 / 63)
    assert scores["d"] == pytest.approx(1 / 63)
    assert scores["a"] > scores["c"]  # 两路命中 > 单路命中
    empty = rrf_fuse([[], ["x"]], k=60)
    assert empty == {"x": pytest.approx(1 / 61)}


# ---------- 3. BM25 jieba 中文检索命中（独立语料） ----------


def test_bm25_jieba_chinese_hit() -> None:
    store.reset_bm25_cache()
    texts = [
        "高血压定义为收缩压≥140mmHg和（或）舒张压≥90mmHg。",
        "二甲双胍是2型糖尿病的一线用药，常用剂量500mg每日两次。",
        "阿司匹林通过抑制血小板聚集发挥抗血栓作用。",
        "慢性阻塞性肺疾病以持续气流受限为特征。",
    ]
    from app.rag.embeddings import _stub_embed

    rows = [
        {
            "chunk_id": f"bm25ut-{i}",
            "kb_document_id": "bm25ut-doc",
            "idx": i,
            "text": t,
            "source_name": "单测.txt",
            "vector": v,
        }
        for i, (t, v) in enumerate(zip(texts, _stub_embed(texts)))
    ]
    store.upsert_chunks(rows, dim=32)
    assert store.rebuild_bm25() >= 4
    hits = store.bm25_search("二甲双胍 剂量", limit=3)
    assert hits and hits[0].metadata["chunk_id"] == "bm25ut-1"
    hits = store.bm25_search("高血压 收缩压", limit=3)
    assert hits and hits[0].metadata["chunk_id"] == "bm25ut-0"
    hits = store.bm25_search("血小板 抗血栓", limit=3)
    assert hits and hits[0].metadata["chunk_id"] == "bm25ut-2"
    store.delete_chunks_for_document("bm25ut-doc")
    store.rebuild_bm25()


# ---------- 4. 检索前置守卫（须在知识库有内容之前运行） ----------


def test_retrieve_guards(client: TestClient, docx_path: Path) -> None:
    with docx_path.open("rb") as f:
        doc = client.post("/api/documents", files={"file": ("medical.docx", f, DOCX_MIME)}).json()
    doc_id = doc["id"]
    # 状态守卫：未分句不能检索
    resp = client.post(f"/api/documents/{doc_id}/retrieve")
    assert resp.status_code == 400 and "解析分句" in resp.json()["detail"]
    # 开关守卫
    assert client.put("/api/settings", json={"retrieve.enabled": False}).status_code == 200
    assert client.post(f"/api/documents/{doc_id}/run").status_code == 200
    resp = client.post(f"/api/documents/{doc_id}/retrieve")
    assert resp.status_code == 400 and "关闭" in resp.json()["detail"]
    assert client.put("/api/settings", json={"retrieve.enabled": True}).status_code == 200
    # 知识库为空守卫（本模块此时尚未上传知识库文档）
    if store.count_chunks() == 0:
        resp = client.post(f"/api/documents/{doc_id}/retrieve")
        assert resp.status_code == 400 and "知识库为空" in resp.json()["detail"]
    assert client.delete(f"/api/documents/{doc_id}").status_code == 200


# ---------- 5. 知识库索引：上传→增量跳过→删除同步 ----------


def test_kb_index_incremental_and_delete(client: TestClient) -> None:
    row1 = _upload_kb(client, "指南.txt", KB_TEXT.encode("utf-8"))
    assert row1["chunk_count"] >= 2
    assert row1["content_hash"]
    chunks_before = store.count_chunks()
    assert chunks_before == row1["chunk_count"]

    # 重新索引：内容未变 → 增量跳过（job 事件含 skipped），chunk 数不变
    reindex = client.post(f"/api/kb/documents/{row1['id']}/reindex")
    assert reindex.status_code == 200, reindex.text
    job_id = reindex.json()["job_id"]
    # SSE 阻塞读取直到 job 结束（规避 reindex 线程与状态轮询的竞态）
    sse = client.get(f"/api/jobs/{job_id}/events")
    assert sse.status_code == 200
    assert "event: skipped" in sse.text and "event: done" in sse.text
    row1_after = _wait_kb_status(client, row1["id"])
    assert row1_after["status"] == "indexed"
    assert row1_after["chunk_count"] == row1["chunk_count"]
    assert store.count_chunks() == chunks_before  # 无重复 chunk

    # 第二个文档上传后删除：chunks 与 BM25 同步清除
    row2 = _upload_kb(client, "补充.txt", ("青霉素使用前必须皮试。" * 30).encode("utf-8"))
    assert store.count_chunks() == chunks_before + row2["chunk_count"]
    hits = store.bm25_search("青霉素 皮试", limit=3)
    assert hits and hits[0].metadata["kb_document_id"] == row2["id"]
    assert client.delete(f"/api/kb/documents/{row2['id']}").status_code == 200
    assert store.count_chunks() == chunks_before
    remaining = store.all_chunks()
    assert all(c["kb_document_id"] != row2["id"] for c in remaining)
    docs = client.get("/api/kb/documents").json()
    assert [d["id"] for d in docs] == [row1["id"]]

    # 收尾：清空知识库，避免污染后续检索全流程测试的 doc_name 断言
    assert client.delete(f"/api/kb/documents/{row1['id']}").status_code == 200
    assert store.count_chunks() == 0


# ---------- 6. 检索全流程：重写 + 3+3 + 跳过规则 + 降级 + SSE ----------


def test_retrieve_full_flow(client: TestClient, docx_path: Path, monkeypatch) -> None:
    row1 = _upload_kb(client, "指南-全流程.txt", KB_TEXT.encode("utf-8"))

    with docx_path.open("rb") as f:
        doc = client.post("/api/documents", files={"file": ("medical2.docx", f, DOCX_MIME)}).json()
    doc_id = doc["id"]
    assert client.post(f"/api/documents/{doc_id}/run").status_code == 200

    # mock LLM：固定返回 3 个重写问题（原句由 rewrite_queries 强制置首）
    def fake_chat_json(system, user, schema_hint=None):
        return {
            "questions": [
                "高血压的诊断标准是什么？",
                "2型糖尿病的一线治疗药物有哪些？",
                "成人血压正常值范围是多少？",
            ]
        }

    monkeypatch.setattr("app.rag.retrieve.chat_json", fake_chat_json)
    resp = client.post(f"/api/documents/{doc_id}/retrieve")
    assert resp.status_code == 200, resp.text
    result = resp.json()
    assert result["status"] == "retrieved"
    assert result["sentences"] > 0
    assert result["evidence"] > 0
    assert result["rewritten"] == result["sentences"]  # mock 下全部重写成功

    # SSE 回放：start / progress / done 完整
    sse = client.get(f"/api/jobs/{result['job_id']}/events")
    assert sse.status_code == 200
    assert "event: start" in sse.text
    assert "event: progress" in sse.text
    assert "event: done" in sse.text

    # 证据结构：跳过规则、queries 溯源、3+3 截断与去重
    data = client.get(f"/api/documents/{doc_id}/evidence").json()
    assert data["status"] == "retrieved"
    sentences = [s for b in data["blocks"] for s in b["sentences"]]
    assert sentences

    normal = [
        s
        for b in data["blocks"]
        if not b["is_reference"]
        for s in b["sentences"]
        if not s["skipped"]
    ]
    assert normal, "应存在待审校的正常句子"
    for s in normal:
        assert len(s["queries"]) == 4  # 原句 + mock 3 问
        assert s["queries"][0]["text"] == s["text"]  # 原句强制第 1 条
        ev = s["evidence"]
        assert len(ev) <= 6  # 3+3 去重后 ≤6
        assert len({e["chunk_text"] for e in ev}) == len(ev)  # 按 chunk 去重无重复
        assert sum(1 for e in ev if e["source"] == "vector") <= 3
        assert sum(1 for e in ev if e["source"] == "keyword") <= 3
        for e in ev:
            assert e["source"] in ("vector", "keyword")
            assert e["doc_name"] == "指南-全流程.txt"
            assert 1 <= e["rank"] <= 3
            assert e["score"] > 0
            assert e["chunk_text"]

    # 占位符句与参考文献块被跳过：无 queries / evidence
    skipped = [s for b in data["blocks"] for s in b["sentences"] if s["skipped"]]
    assert any(s["text"].startswith("[{表格不予审校_") for s in skipped)
    assert any(b["is_reference"] for b in data["blocks"])
    for s in skipped:
        assert s["queries"] == [] and s["evidence"] == []

    # LLM 未配置 → 降级为只用原句检索（不报错）
    def raise_not_configured(*args, **kwargs):
        raise LLMNotConfiguredError("未配置 LLM API Key")

    monkeypatch.setattr("app.rag.retrieve.chat_json", raise_not_configured)
    resp2 = client.post(f"/api/documents/{doc_id}/retrieve")
    assert resp2.status_code == 200, resp2.text
    assert resp2.json()["rewritten"] == 0
    data2 = client.get(f"/api/documents/{doc_id}/evidence").json()
    normal2 = [
        s for b in data2["blocks"] if not b["is_reference"] for s in b["sentences"] if not s["skipped"]
    ]
    for s in normal2:
        assert [q["text"] for q in s["queries"]] == [s["text"]]  # 仅原句
        assert len(s["evidence"]) <= 6  # 降级后仍能检索出证据

    assert client.delete(f"/api/documents/{doc_id}").status_code == 200
    assert client.delete(f"/api/kb/documents/{row1['id']}").status_code == 200


def test_job_events_404(client: TestClient) -> None:
    assert client.get("/api/jobs/nonexistent-job/events").status_code == 404
