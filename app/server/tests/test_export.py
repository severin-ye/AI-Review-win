"""M5 导出 E2E：mock LLM 全链路 upload→run→retrieve→review→decisions→export。

- embedding 走 stub provider、LLM 一律 mock（retrieve / review 两处 chat_json），
  零真实 API/模型/网络依赖；环境隔离同 M2-M4（AI_REVIEW_DATA_DIR + AI_REVIEW_SEGMENTER=rule）。
- 断言双版本 docx 语义：清洁版直接替换、留痕版 ｛～原文:x AI/用户:y～｝、表格还原、
  占位符不残留、首行缩进、Heading 标题、adopted 计数、锚点缺失降级 warning。
"""
import json
import os
import re
import tempfile
import time
from pathlib import Path

_tmp = tempfile.mkdtemp(prefix="ai-review-test-export-")
os.environ.setdefault("AI_REVIEW_DATA_DIR", _tmp)
os.environ.setdefault("AI_REVIEW_SEGMENTER", "rule")

import pytest  # noqa: E402
from docx import Document as DocxDocument  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402
from sqlmodel import Session, select  # noqa: E402

from app.core.config import get_settings  # noqa: E402
from app.core.db import engine  # noqa: E402
from app.main import app  # noqa: E402
from app.models import Block, Correction, Document, Sentence  # noqa: E402

DATA_DIR = get_settings().data_dir
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

KB_TEXT = """中国高血压防治指南：高血压定义为在未使用降压药物的情况下，非同日三次测量诊室血压，收缩压≥140mmHg和（或）舒张压≥90mmHg。根据血压升高水平，将高血压分为1级、2级和3级。降压治疗的根本目标是降低心、脑、肾与血管并发症和死亡的总危险。

2型糖尿病防治要点：二甲双胍是2型糖尿病的一线首选用药，无禁忌证者应一直保留在治疗方案中。常用起始剂量500mg每日两次，随餐服用可减少胃肠道反应。糖化血红蛋白控制目标一般为低于7.0%，需根据患者年龄与并发症个体化调整。

阿司匹林肠溶片说明：阿司匹林通过不可逆抑制环氧化酶，阻断血栓素A2生成，从而抑制血小板聚集，广泛用于动脉粥样硬化性心血管疾病的二级预防。常用剂量为每日75至100mg，长期服用需注意消化道出血风险，禁用于活动性溃疡患者。

慢性阻塞性肺疾病诊治：慢阻肺以持续气流受限为特征，肺功能检查吸入支气管扩张剂后FEV1/FVC小于0.70可确诊。稳定期首选长效支气管扩张剂，急性加重期可短程使用全身糖皮质激素并给予控制性氧疗，目标血氧饱和度88%至92%。
"""

# 四句正文各有约定俗成的 correction（见 fake_review_chat_json）
PARA_ACCEPT = "患者男性，六十二岁，确诊高血压三年，目前规律服用氨氯地平控制血压。"
PARA_CUSTOM = "入院查体血压150/95mmHg，心肺听诊未见明显异常。"
PARA_REJECT = "患者要求出院，嘱其继续服药并定期复查血压变化情况。"


def _build_export_docx(path: Path) -> None:
    """含 2 表格（首表按审查意见表移除，第二表 → 占位符）+ 章节 + 参考文献的测试文档。"""
    doc = DocxDocument()
    doc.add_heading("一、病例资料", level=1)
    doc.add_paragraph(PARA_ACCEPT)
    doc.add_paragraph(PARA_CUSTOM)
    table1 = doc.add_table(rows=2, cols=2)  # 首个表格：审查意见表（has_review_table=Y 时移除）
    table1.cell(0, 0).text = "审查意见"
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "检查项目"
    table2.cell(0, 1).text = "结果"
    doc.add_paragraph(PARA_REJECT)
    doc.add_heading("参考文献", level=2)
    doc.add_paragraph("[1]: 中国高血压防治指南修订委员会. 中华心血管病杂志, 2019.")
    doc.save(str(path))


@pytest.fixture(scope="module")
def client():
    with TestClient(app) as c:
        assert c.put(
            "/api/settings",
            json={
                "embedding.provider": "stub",
                "retrieve.enabled": True,
                "docx.has_review_table": "Y",
                "output.dir": "",
                "docx.first_line_indent": 0.5,
                "llm.base_url": "http://127.0.0.1:9/v1",  # 永不可达，chat_json 均被 mock
                "llm.api_key": "sk-export-test",
                "llm.model": "test-model",
            },
        ).status_code == 200
        yield c


@pytest.fixture(scope="module")
def docx_path() -> Path:
    path = Path(_tmp) / "export-sample.docx"
    _build_export_docx(path)
    return path


def _wait_job(client: TestClient, job_id: str, timeout: float = 15.0) -> str:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        status = client.get(f"/api/jobs/{job_id}").json()["status"]
        if status in ("done", "error"):
            return status
        time.sleep(0.1)
    raise AssertionError(f"job {job_id} 超时未完成")


def _wait_kb_indexed(client: TestClient, kb_id: str, timeout: float = 90) -> None:
    deadline = time.time() + timeout
    while time.time() < deadline:
        docs = client.get("/api/kb/documents").json()
        row = next((d for d in docs if d["id"] == kb_id), None)
        if row and row["status"] == "indexed":
            return
        if row and row["status"] == "failed":
            raise AssertionError(f"知识库索引失败: {row}")
        time.sleep(0.5)
    raise TimeoutError("知识库索引超时")


def fake_review_chat_json(system, user, schema_hint=None):
    """按 user prompt 中 [S] 编号行的内容生成四条约定 correction：

    - 含「氨氯地平」句：accepted 候补（锚点正常）+ 一条 original 非逐字摘录的坏锚点条目；
    - 含「150/95mmHg」句：custom 候补；
    - 含「继续服药」句：rejected 候补。
    """
    corrections = []
    for line in user.splitlines():
        match = re.match(r"^\[S(\d+)\]\s*(.*)$", line)
        if not match:
            continue
        num, text = int(match.group(1)), match.group(2)
        if "氨氯地平" in text:
            corrections.append(
                {
                    "sentence_id": num,
                    "original": "氨氯地平",
                    "suggestion": "苯磺酸氨氯地平",
                    "error_type": "术语错误",
                    "severity": "medium",
                    "evidence_ids": [],
                    "explanation": "药品通用名应完整。",
                }
            )
            corrections.append(
                {
                    "sentence_id": num,
                    "original": "硝苯地平控释片",  # 非逐字摘录：导出时应降级 warning
                    "suggestion": "硝苯地平",
                    "error_type": "术语错误",
                    "severity": "low",
                    "evidence_ids": [],
                    "explanation": "锚点缺失用例。",
                }
            )
        if "150/95mmHg" in text:
            corrections.append(
                {
                    "sentence_id": num,
                    "original": "150/95mmHg",
                    "suggestion": "140/90mmHg",
                    "error_type": "事实错误",
                    "severity": "high",
                    "evidence_ids": [1],
                    "explanation": "数值待人工核对 [E1]。",
                }
            )
        if "继续服药" in text:
            corrections.append(
                {
                    "sentence_id": num,
                    "original": "定期复查血压",
                    "suggestion": "每四周复查血压",
                    "error_type": "格式错误",
                    "severity": "low",
                    "evidence_ids": [],
                    "explanation": "随访频率建议明确。",
                }
            )
    return {"corrections": corrections}


def _doc_status(doc_id: str) -> str:
    with Session(engine) as session:
        return session.get(Document, doc_id).status


def _paragraph_texts(doc: DocxDocument) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def test_export_full_flow(client: TestClient, docx_path: Path, monkeypatch) -> None:
    # ---------- 知识库（retrieve 前置） ----------
    kb = client.post(
        "/api/kb/documents", files={"file": ("指南-export.txt", KB_TEXT.encode("utf-8"), "text/plain")}
    )
    assert kb.status_code == 201, kb.text
    kb_id = kb.json()["id"]
    _wait_kb_indexed(client, kb_id)

    # ---------- upload → run（ingest + segment） ----------
    with docx_path.open("rb") as f:
        up = client.post("/api/documents", files={"file": ("export-sample.docx", f, DOCX_MIME)})
    assert up.status_code == 201, up.text
    doc_id = up.json()["id"]
    run = client.post(f"/api/documents/{doc_id}/run")
    assert run.status_code == 200 and run.json()["status"] == "segmented", run.text
    parsed = client.get(f"/api/documents/{doc_id}/parsed")
    assert "[{表格不予审校_1}]" in parsed.text

    # ---------- retrieve（stub embedding + mock 查询重写） ----------
    monkeypatch.setattr(
        "app.rag.retrieve.chat_json",
        lambda system, user, schema_hint=None: {"questions": ["高血压的诊断标准是什么？"]},
    )
    retr = client.post(f"/api/documents/{doc_id}/retrieve")
    assert retr.status_code == 200 and retr.json()["status"] == "retrieved", retr.text

    # ---------- review（mock LLM 出四条约定 correction） ----------
    monkeypatch.setattr("app.pipeline.review.chat_json", fake_review_chat_json)
    rev = client.post(f"/api/documents/{doc_id}/review")
    assert rev.status_code == 200, rev.text
    assert _wait_job(client, rev.json()["job_id"]) == "done"

    detail = client.get(f"/api/documents/{doc_id}/detail").json()
    all_corrections = [
        (s, c)
        for b in detail["blocks"]
        for s in b["sentences"]
        for c in s["corrections"]
    ]
    assert len(all_corrections) == 4, f"应产出 4 条 correction，实际 {len(all_corrections)}"

    def find_correction(**kw):
        for _, c in all_corrections:
            if all(c[k] == v for k, v in kw.items()):
                return c
        raise AssertionError(f"未找到 correction: {kw}")

    # ---------- decisions：accept / custom / reject 各一 + 坏锚点 accept ----------
    c_accept = find_correction(suggestion="苯磺酸氨氯地平")
    c_bad_anchor = find_correction(original="硝苯地平控释片")
    c_custom = find_correction(original="150/95mmHg")
    c_reject = find_correction(suggestion="每四周复查血压")

    assert client.post(
        f"/api/corrections/{c_accept['id']}/decision", json={"decision": "accepted"}
    ).status_code == 200
    assert client.post(
        f"/api/corrections/{c_bad_anchor['id']}/decision", json={"decision": "accepted"}
    ).status_code == 200
    resp = client.post(
        f"/api/corrections/{c_custom['id']}/decision",
        json={"decision": "custom", "custom_text": "138/88mmHg"},
    )
    assert resp.status_code == 200, resp.text
    resp = client.post(
        f"/api/corrections/{c_reject['id']}/decision", json={"decision": "rejected"}
    )
    assert resp.status_code == 200
    assert resp.json()["document_status"] == "manual_done"

    # ---------- export ----------
    exp = client.post(f"/api/documents/{doc_id}/export")
    assert exp.status_code == 200, exp.text
    result = exp.json()
    assert result["status"] == "done"
    assert result["adopted"] == 3  # accepted×2（含坏锚点）+ custom×1；rejected 不计
    assert len(result["warnings"]) == 1  # 坏锚点降级
    assert "硝苯地平控释片" in result["warnings"][0]

    clean_path = Path(result["clean_path"])
    marked_path = Path(result["marked_path"])
    assert clean_path.name == "export-sample_审校修订1_.docx"
    assert marked_path.name == "export-sample_审校修订2_.docx"
    assert clean_path.exists() and marked_path.exists()
    # 默认输出目录 projects/<doc_id>/exports/
    assert clean_path.parent == DATA_DIR / "projects" / doc_id / "exports"

    # 文档状态与 exports_json 落库
    assert _doc_status(doc_id) == "done"
    with Session(engine) as session:
        row = session.get(Document, doc_id)
        payload = json.loads(row.exports_json)
        assert payload["adopted"] == 3
        assert payload["clean"] == str(clean_path) and payload["marked"] == str(marked_path)

    # job 事件：start / progress / warning / done 齐备（SSE 回放）
    sse = client.get(f"/api/jobs/{result['job_id']}/events")
    assert sse.status_code == 200
    assert "event: start" in sse.text and "event: progress" in sse.text
    assert "event: warning" in sse.text  # 坏锚点降级事件
    assert "event: done" in sse.text

    # ---------- 清洁版内容断言 ----------
    clean = DocxDocument(str(clean_path))
    clean_text = _paragraph_texts(clean)
    assert "苯磺酸氨氯地平" in clean_text  # accepted → suggestion
    assert "138/88mmHg" in clean_text and "140/90mmHg" not in clean_text  # custom → custom_text
    assert "定期复查血压" in clean_text and "每四周复查血压" not in clean_text  # rejected 保留原文
    assert "硝苯地平" not in clean_text  # 坏锚点：保留原文
    assert "｛～" not in clean_text  # 清洁版无任何留痕标记
    assert "[{表格不予审校" not in clean_text  # 占位符不残留
    assert len(clean.tables) == 1  # tables.docx 第 1 个表格被还原
    assert clean.tables[0].cell(0, 0).text == "检查项目"
    # 章节标题用 Heading 样式
    heading = next(p for p in clean.paragraphs if p.text == "一、病例资料")
    assert heading.style.name.startswith(("Heading", "标题"))
    # 全文段落首行缩进 0.5 英寸
    assert clean.paragraphs, "导出文档应至少有一个段落"
    for p in clean.paragraphs:
        assert p.paragraph_format.first_line_indent == 457200, (
            f"段落 {p.text[:20]!r} 首行缩进缺失（Emu 457200 = 0.5in）"
        )
    # 参考文献块原文保留
    assert "[1]: 中国高血压防治指南修订委员会. 中华心血管病杂志, 2019." in clean_text

    # ---------- 留痕版内容断言 ----------
    marked = DocxDocument(str(marked_path))
    marked_text = _paragraph_texts(marked)
    assert "｛～原文:氨氯地平 AI:苯磺酸氨氯地平～｝" in marked_text
    assert "｛～原文:150/95mmHg 用户:138/88mmHg～｝" in marked_text
    assert "｛～原文:硝苯地平控释片" not in marked_text  # 坏锚点不留痕
    assert "定期复查血压" in marked_text and "｛～原文:定期复查血压" not in marked_text
    assert "[{表格不予审校" not in marked_text
    assert len(marked.tables) == 1

    # ---------- exports 列表与下载 ----------
    listing = client.get(f"/api/documents/{doc_id}/exports")
    assert listing.status_code == 200
    body = listing.json()
    assert body["adopted"] == 3
    assert {e["kind"] for e in body["exports"]} == {1, 2}
    assert all(e["exists"] and e["size"] > 0 for e in body["exports"])

    dl = client.get(f"/api/documents/{doc_id}/exports/1")
    assert dl.status_code == 200 and dl.content[:2] == b"PK"  # docx = zip
    dl2 = client.get(f"/api/documents/{doc_id}/exports/2")
    assert dl2.status_code == 200 and dl2.content[:2] == b"PK"
    assert client.get(f"/api/documents/{doc_id}/exports/3").status_code == 404

    # ---------- 重导出幂等（done 状态允许重跑） ----------
    exp2 = client.post(f"/api/documents/{doc_id}/export")
    assert exp2.status_code == 200 and exp2.json()["adopted"] == 3

    # ---------- 清理 ----------
    assert client.delete(f"/api/documents/{doc_id}").status_code == 200
    assert client.delete(f"/api/kb/documents/{kb_id}").status_code == 200


def test_export_guards(client: TestClient, docx_path: Path) -> None:
    assert client.post("/api/documents/nonexistent/export").status_code == 404

    with docx_path.open("rb") as f:
        up = client.post("/api/documents", files={"file": ("guard.docx", f, DOCX_MIME)})
    doc_id = up.json()["id"]
    # uploaded 状态不可导出
    resp = client.post(f"/api/documents/{doc_id}/export")
    assert resp.status_code == 400 and "不能导出" in resp.json()["detail"]
    # 未导出过时 exports 为空、下载 404
    assert client.get(f"/api/documents/{doc_id}/exports").json()["exports"] == []
    assert client.get(f"/api/documents/{doc_id}/exports/1").status_code == 404
    assert client.delete(f"/api/documents/{doc_id}").status_code == 200


def test_export_without_adopted_corrections(client: TestClient, docx_path: Path, monkeypatch) -> None:
    """零 accepted/custom 决定：仍导出（等同原文），adopted=0。"""
    with docx_path.open("rb") as f:
        up = client.post("/api/documents", files={"file": ("plain.docx", f, DOCX_MIME)})
    doc_id = up.json()["id"]
    assert client.post(f"/api/documents/{doc_id}/run").status_code == 200

    # 直接落种一条 correction 并 reject（不经 review，直奔人工决定）
    with Session(engine) as session:
        doc = session.get(Document, doc_id)
        doc.status = "pending_manual"
        session.add(doc)
        block = session.exec(select(Block).where(Block.document_id == doc_id)).first()
        sent = session.exec(
            select(Sentence).where(Sentence.block_id == block.id, Sentence.idx > 0)
        ).first()
        corr = Correction(
            sentence_id=sent.id,
            original=sent.text[:4],
            suggestion="任意建议",
            error_type="格式错误",
            severity="low",
        )
        session.add(corr)
        session.commit()
        session.refresh(corr)
        corr_id = corr.id
    resp = client.post(f"/api/corrections/{corr_id}/decision", json={"decision": "rejected"})
    assert resp.status_code == 200
    assert resp.json()["document_status"] == "manual_done"

    exp = client.post(f"/api/documents/{doc_id}/export")
    assert exp.status_code == 200, exp.text
    result = exp.json()
    assert result["adopted"] == 0
    clean = DocxDocument(result["clean_path"])
    text = _paragraph_texts(clean)
    assert "任意建议" not in text and PARA_ACCEPT in text
    assert client.delete(f"/api/documents/{doc_id}").status_code == 200
