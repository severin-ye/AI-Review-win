"""知识库文档索引：加载 → LangChain 切块 → 稳定 ID → 增量嵌入 → LanceDB + BM25。

- 切块：LangChain RecursiveCharacterTextSplitter（chunk 500 / overlap 50，中文分隔符优先）。
- 稳定 chunk_id = sha256(content_hash:idx)[:16]：同一文档重切块结果 ID 稳定，幂等 upsert。
- 增量索引（修复旧版"永不重建"bug）：content_hash 未变且已 indexed → 跳过；
  reindex 按新 hash 重切重嵌；删除文档同步删 chunks 并重建 BM25。
- 进度通过 emit(event, data) 回调写 job_events（API 层接线，SSE 推送）。
"""
from __future__ import annotations

import csv
import hashlib
from pathlib import Path
from typing import Any, Callable

from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlmodel import Session

from app.core.db import engine
from app.models import KbDocument
from app.rag import store
from app.rag.embeddings import EmbeddingProvider

CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
EMBED_BATCH = 32

# 中文文本优先按段落/句读切（设计文档 §5.3：chunk 500 字 / overlap 50 沿用旧版）
_SPLITTER = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    separators=["\n\n", "\n", "。", "；", "，", " ", ""],
)

Emit = Callable[[str, dict], None]

SUPPORTED_SUFFIXES = {".pdf", ".txt", ".csv", ".docx"}


def kb_files_dir() -> Path:
    directory = store.kb_dir() / "files"
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def kb_file_path(kb_document_id: str, suffix: str) -> Path:
    return kb_files_dir() / f"{kb_document_id}{suffix.lower()}"


def content_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def chunk_id(doc_hash: str, idx: int) -> str:
    return hashlib.sha256(f"{doc_hash}:{idx}".encode("utf-8")).hexdigest()[:16]


def load_kb_file(path: Path) -> str:
    """按扩展名加载知识库文档为纯文本：PDF(pypdf) / TXT / CSV / DOCX(python-docx)。"""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".csv":
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
            rows = list(csv.reader(f))
        return "\n".join("，".join(cell.strip() for cell in row if cell.strip()) for row in rows)
    if suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:  # 表格按行并入文本，知识库不保留结构
            for row in table.rows:
                line = "，".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if line:
                    parts.append(line)
        return "\n".join(parts)
    raise ValueError(f"不支持的文件类型: {suffix}（支持 pdf/txt/csv/docx）")


def split_chunks(text: str) -> list[str]:
    """RecursiveCharacterTextSplitter 切块，过滤空白块。"""
    return [c.strip() for c in _SPLITTER.split_text(text) if c.strip()]


def _noop_emit(_event: str, _data: dict) -> None:
    pass


def index_kb_document(kb_document_id: str, emit: Emit = _noop_emit) -> dict[str, Any]:
    """索引单个知识库文档（幂等、增量）。

    返回 {"status": "indexed"|"skipped", "chunks": n}；失败置 kb_documents.status=failed 并抛出。
    """
    with Session(engine) as session:
        row = session.get(KbDocument, kb_document_id)
        if row is None:
            raise KeyError(f"知识库文档不存在: {kb_document_id}")
        filename = row.filename
        old_hash = row.content_hash
        already_indexed = row.status == "indexed" and row.chunk_count > 0
        row.status = "indexing"
        session.add(row)
        session.commit()

    suffix = Path(filename).suffix.lower()
    path = kb_file_path(kb_document_id, suffix)
    if not path.exists():
        raise FileNotFoundError(f"知识库文件缺失: {path}")

    try:
        emit("start", {"kb_document_id": kb_document_id, "filename": filename})
        doc_hash = content_hash(path.read_bytes())

        # 增量索引：内容未变且已成功索引过 → 跳过（修复旧版"永不重建/重复重建"问题）
        if already_indexed and old_hash == doc_hash:
            with Session(engine) as session:
                row = session.get(KbDocument, kb_document_id)
                row.status = "indexed"
                session.add(row)
                session.commit()
                skipped_chunks = row.chunk_count
            emit("skipped", {"reason": "内容未变更，增量跳过", "content_hash": doc_hash[:12]})
            return {"status": "skipped", "chunks": skipped_chunks}

        text = load_kb_file(path)
        emit("loaded", {"chars": len(text)})
        chunks = split_chunks(text)
        emit("chunked", {"chunks": len(chunks), "chunk_size": CHUNK_SIZE, "overlap": CHUNK_OVERLAP})
        if not chunks:
            raise ValueError("文档未切出有效文本块（内容为空或无法解析）")

        provider = EmbeddingProvider.get()
        rows: list[dict[str, Any]] = []
        for start in range(0, len(chunks), EMBED_BATCH):
            batch = chunks[start : start + EMBED_BATCH]
            vectors = provider.embed(batch)
            for offset, (chunk_text, vector) in enumerate(zip(batch, vectors)):
                idx = start + offset
                rows.append(
                    {
                        "chunk_id": chunk_id(doc_hash, idx),
                        "kb_document_id": kb_document_id,
                        "idx": idx,
                        "text": chunk_text,
                        "source_name": filename,
                        "vector": vector,
                    }
                )
            emit("embedding", {"done": min(start + EMBED_BATCH, len(chunks)), "total": len(chunks)})

        store.delete_chunks_for_document(kb_document_id)  # 重索引先清旧 chunks
        store.upsert_chunks(rows, dim=len(rows[0]["vector"]))
        bm25_count = store.rebuild_bm25()
        emit("bm25", {"indexed_chunks": bm25_count})

        with Session(engine) as session:
            row = session.get(KbDocument, kb_document_id)
            row.content_hash = doc_hash
            row.status = "indexed"
            row.chunk_count = len(chunks)
            session.add(row)
            session.commit()
        emit("done", {"chunks": len(chunks), "content_hash": doc_hash[:12]})
        return {"status": "indexed", "chunks": len(chunks)}
    except Exception as exc:
        with Session(engine) as session:
            row = session.get(KbDocument, kb_document_id)
            if row is not None:
                row.status = "failed"
                session.add(row)
                session.commit()
        emit("error", {"message": str(exc)})
        raise


def delete_kb_document(kb_document_id: str) -> None:
    """删除知识库文档：LanceDB chunks + BM25 重建 + 落盘文件 + DB 行。"""
    store.delete_chunks_for_document(kb_document_id)
    store.rebuild_bm25()
    for path in kb_files_dir().glob(f"{kb_document_id}.*"):
        path.unlink(missing_ok=True)
    with Session(engine) as session:
        row = session.get(KbDocument, kb_document_id)
        if row is not None:
            session.delete(row)
            session.commit()
