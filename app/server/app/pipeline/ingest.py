"""① ingest：docx 解析（表格占位 + 转 Markdown）。

移植旧项目 src/utils/table_utils.py 验证过的思路（只读参考，不改动旧文件）：
a. settings 中 docx.has_review_table == 'Y' 时移除第一个表格（审查意见表）
b. 其余表格按顺序抽取另存 tables.docx
c. 原文档中表格原位替换为占位符段落 [{表格不予审校_N}]（N 从 1 递增），存 no_table.docx
d. no_table.docx → parsed.md：Heading 1-3 → #/##/###，其余按段落输出，保留空行分段
"""
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn

from app.core.user_settings import has_review_table
from app.pipeline.common import project_dir, set_document_status

PLACEHOLDER_RE = re.compile(r"\[\{表格不予审校_(\d+)\}\]")
_HEADING_STYLE_RE = re.compile(r"(?i)^\s*(heading|标题)\s*([1-3])\s*$")


def _remove_first_table(doc: DocxDocument) -> bool:
    """移除首个表格（审查意见表）。返回是否有表格被移除。"""
    if not doc.tables:
        return False
    first_tbl = doc.tables[0]._element
    doc.element.body.remove(first_tbl)
    return True


def _extract_tables(doc: DocxDocument, dest_path: Path) -> int:
    """把文档中所有（剩余）表格按顺序抽取到独立 docx。"""
    dst = DocxDocument()
    body = dst._element.body
    sectPr = body.find(qn("w:sectPr"))
    for table in doc.tables:
        new_tbl = deepcopy(table._tbl)
        if sectPr is not None:
            sectPr.addprevious(new_tbl)  # 保持 sectPr 位于文末，避免损坏文档结构
        else:
            body.append(new_tbl)
        dst.add_paragraph()
    dst.save(str(dest_path))
    return len(doc.tables)


def _replace_tables_with_placeholders(doc: DocxDocument, dest_path: Path) -> int:
    """表格原位替换为 [{表格不予审校_N}] 占位符段落，返回占位符数量。"""
    body = doc.element.body
    counter = 1
    for element in list(body):
        if element.tag == qn("w:tbl"):
            index = body.index(element)
            body.remove(element)
            placeholder = doc.add_paragraph(f"[{{表格不予审校_{counter}}}]")
            body.insert(index, placeholder._element)
            counter += 1
    doc.save(str(dest_path))
    return counter - 1


def _heading_level(para) -> int | None:
    """识别 Heading 1-3 / 标题 1-3 段落样式。"""
    try:
        name = para.style.name or ""
    except Exception:  # 样式缺失时按普通段落处理
        return None
    match = _HEADING_STYLE_RE.match(name)
    return int(match.group(2)) if match else None


def _to_markdown(doc: DocxDocument) -> str:
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.rstrip()
        level = _heading_level(para)
        if level is not None and text.strip():
            lines.append("#" * level + " " + text.strip())
        else:
            # 普通段落与表格占位符原样输出；空段落保留为空行以维持分段
            lines.append(text)
    return "\n".join(lines).strip("\n") + "\n"


def ingest_document(doc_id: str) -> dict:
    """解析 projects/<doc_id>/original.docx，产出 tables.docx / no_table.docx / parsed.md。

    成功：状态 uploaded → parsed；失败：记 error 并置 failed 后重新抛出。
    """
    pdir = project_dir(doc_id)
    original = pdir / "original.docx"
    if not original.exists():
        raise FileNotFoundError(f"未找到上传文件: {original}")

    try:
        doc = DocxDocument(str(original))
        removed_review_table = has_review_table() and _remove_first_table(doc)

        tables_path = pdir / "tables.docx"
        no_table_path = pdir / "no_table.docx"
        parsed_path = pdir / "parsed.md"

        table_count = _extract_tables(doc, tables_path)
        placeholder_count = _replace_tables_with_placeholders(doc, no_table_path)

        md_doc = DocxDocument(str(no_table_path))
        parsed_path.write_text(_to_markdown(md_doc), encoding="utf-8")

        set_document_status(doc_id, "parsed")
        return {
            "removed_review_table": removed_review_table,
            "tables": table_count,
            "placeholders": placeholder_count,
            "parsed_md": str(parsed_path),
        }
    except Exception as exc:
        set_document_status(doc_id, "failed", str(exc))
        raise
