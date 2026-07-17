"""M2 流水线测试：docx 解析（表格占位）+ 分句分块入库，走 TestClient 全链路。

测试通过 AI_REVIEW_SEGMENTER=rule 强制正则分句回退路径（不依赖 SaT 模型下载）。
"""
import os
import re
import tempfile
from pathlib import Path

# 在导入 app 之前隔离数据目录与分句器后端，避免污染 dev 数据 / 触发模型下载
_tmp = tempfile.mkdtemp(prefix="ai-review-test-pipeline-")
os.environ["AI_REVIEW_DATA_DIR"] = _tmp
os.environ["AI_REVIEW_SEGMENTER"] = "rule"

import pytest  # noqa: E402
from docx import Document as DocxDocument  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

from app.core.config import get_settings  # noqa: E402
from app.main import app  # noqa: E402
from app.pipeline.segment import SentenceSplitter, merge_short_sentences  # noqa: E402

# 注意：与其他测试模块同跑时，引擎绑定的是首个导入模块设置的数据目录；
# 因此文件路径断言一律以 get_settings().data_dir 为准，不用本模块的 _tmp。
DATA_DIR = get_settings().data_dir

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PLACEHOLDER_RE = re.compile(r"^\[\{表格不予审校_\d+\}\]$")
CHAPTER_RE = re.compile(r"^\s*(?:[一二三四五六七八九十百]+、|（[一二三四五六七八九十]+）)")


def _build_test_docx(path: Path) -> None:
    """含 2 个表格 + 「一、二、」章节 + 短句碎片 + [N]: 参考文献行的测试文档。"""
    doc = DocxDocument()
    doc.add_heading("一、病例资料", level=1)
    doc.add_paragraph("患者男性，五十六岁，因反复胸闷气短三个月入院治疗。")
    table1 = doc.add_table(rows=2, cols=2)  # 首个表格：模拟审查意见表（默认被移除）
    table1.cell(0, 0).text = "审查意见"
    doc.add_paragraph("入院后完善相关检查，血常规提示白细胞计数轻度升高。")
    doc.add_paragraph("患者恢复良好。嗯。复查胸部CT未见明显异常改变。")
    doc.add_paragraph("二、诊疗经过")
    doc.add_paragraph("入院后给予抗感染、化痰及对症支持治疗十天左右。")
    table2 = doc.add_table(rows=2, cols=2)
    table2.cell(0, 0).text = "检查项目"
    doc.add_heading("参考文献", level=2)
    doc.add_paragraph("[1]: 张三. 中华医学杂志, 2020, 100(1): 1-5.")
    doc.add_paragraph("[2]: 李四. 中国实用内科杂志, 2021, 41(2): 100-103.")
    doc.save(str(path))


@pytest.fixture(scope="module")
def client():
    # with 形式触发 lifespan（启动时 init_db 建表）
    with TestClient(app) as c:
        yield c


@pytest.fixture(scope="module")
def docx_path() -> Path:
    path = Path(_tmp) / "sample.docx"
    _build_test_docx(path)
    return path


def _all_sentences(detail: dict) -> list[tuple[dict, dict]]:
    """返回 (block, sentence) 对，便于全文断言。"""
    return [(b, s) for b in detail["blocks"] for s in b["sentences"]]


def test_rule_splitter_and_merge() -> None:
    splitter = SentenceSplitter.get()
    assert splitter.backend == "rule"
    text = "患者男性，56岁。因反复胸闷、气短3个月入院。版本号3.5不拆。血常规正常；继续观察。"
    sents = splitter.split(text)
    assert "患者男性，56岁。" in sents
    assert "因反复胸闷、气短3个月入院。" in sents
    assert "版本号3.5不拆。" in sents  # 数字间小数点不拆
    assert "血常规正常；" in sents
    # 短句合并：碎片并入前一句，不丢文本
    merged = merge_short_sentences(["这是一个足够长的正常句子。", "嗯。", "后续继续观察治疗直至痊愈出院。"], 10)
    assert merged == ["这是一个足够长的正常句子。嗯。", "后续继续观察治疗直至痊愈出院。"]
    # 开头碎片向后并入下一句
    merged2 = merge_short_sentences(["好。", "这是一个足够长的正常句子。"], 10)
    assert merged2 == ["好。这是一个足够长的正常句子。"]


def test_upload_run_detail_flow(client: TestClient, docx_path: Path) -> None:
    # 1. 上传
    with docx_path.open("rb") as f:
        resp = client.post("/api/documents", files={"file": ("sample.docx", f, DOCX_MIME)})
    assert resp.status_code == 201, resp.text
    doc = resp.json()
    doc_id = doc["id"]
    assert doc["filename"] == "sample.docx"
    assert doc["status"] == "uploaded"

    # 列表可见
    listing = client.get("/api/documents")
    assert listing.status_code == 200
    assert any(d["id"] == doc_id for d in listing.json())

    # 2. 运行流水线（ingest + segment）
    run = client.post(f"/api/documents/{doc_id}/run")
    assert run.status_code == 200, run.text
    result = run.json()
    assert result["status"] == "segmented"
    assert result["segmenter"] == "rule"
    assert result["blocks"] > 0 and result["sentences"] > 0

    # 产物文件：tables.docx / no_table.docx / parsed.md
    project = DATA_DIR / "projects" / doc_id
    assert (project / "tables.docx").exists()
    assert (project / "no_table.docx").exists()
    parsed = client.get(f"/api/documents/{doc_id}/parsed")
    assert parsed.status_code == 200
    # 默认 has_review_table=Y：首个表格被移除，仅剩 1 个占位符
    assert "[{表格不予审校_1}]" in parsed.text
    assert "[{表格不予审校_2}]" not in parsed.text

    # 3. 详情：章节分组 / 占位符 / 参考文献块 / 短句合并
    detail = client.get(f"/api/documents/{doc_id}/detail").json()
    blocks = detail["blocks"]
    assert [b["idx"] for b in blocks] == list(range(len(blocks)))
    chapters = [b["chapter"] for b in blocks]
    assert "一、病例资料" in chapters
    assert "二、诊疗经过" in chapters
    assert "参考文献" in chapters

    # 表格占位符独立成句，出现在「二、诊疗经过」块
    diag_block = next(b for b in blocks if b["chapter"] == "二、诊疗经过")
    assert any(PLACEHOLDER_RE.match(s["text"]) for s in diag_block["sentences"])

    # 参考文献区单独成块
    ref_blocks = [b for b in blocks if b["is_reference"]]
    assert len(ref_blocks) == 1
    ref_texts = [s["text"] for s in ref_blocks[0]["sentences"]]
    assert ref_texts == [
        "[1]: 张三. 中华医学杂志, 2020, 100(1): 1-5.",
        "[2]: 李四. 中国实用内科杂志, 2021, 41(2): 100-103.",
    ]

    # 短句碎片「嗯。」被合并（并入累积的前一句），不单独成句、不丢文本
    all_texts = [s["text"] for _, s in _all_sentences(detail)]
    assert "嗯。" not in all_texts
    assert any(t.startswith("患者恢复良好。嗯。") for t in all_texts)

    # 无短碎片句：<10 字只可能是章节标题 / 占位符 / 参考文献行
    for block, sent in _all_sentences(detail):
        text = sent["text"]
        allowed_short = (
            text == block["chapter"]
            or CHAPTER_RE.match(text) is not None
            or PLACEHOLDER_RE.match(text) is not None
            or block["is_reference"]
        )
        assert len(text) >= 10 or allowed_short, f"发现未合并的短碎片句: {text!r}"

    # 4. 可重复运行：重跑后 blocks/sentences 不翻倍
    rerun = client.post(f"/api/documents/{doc_id}/run")
    assert rerun.status_code == 200
    assert rerun.json()["blocks"] == result["blocks"]
    assert rerun.json()["sentences"] == result["sentences"]

    # 5. has_review_table=N 时保留全部表格 → 2 个占位符
    assert client.put("/api/settings", json={"docx.has_review_table": "N"}).status_code == 200
    rerun2 = client.post(f"/api/documents/{doc_id}/run")
    assert rerun2.status_code == 200
    parsed2 = client.get(f"/api/documents/{doc_id}/parsed")
    assert "[{表格不予审校_1}]" in parsed2.text
    assert "[{表格不予审校_2}]" in parsed2.text

    # 6. 删除：库行与项目目录一并清理
    assert client.delete(f"/api/documents/{doc_id}").status_code == 200
    assert client.get(f"/api/documents/{doc_id}/detail").status_code == 404
    assert not project.exists()


def test_run_missing_document(client: TestClient) -> None:
    assert client.post("/api/documents/nonexistent/run").status_code == 404


def test_upload_rejects_non_docx(client: TestClient) -> None:
    resp = client.post("/api/documents", files={"file": ("a.txt", b"hello", "text/plain")})
    assert resp.status_code == 400
