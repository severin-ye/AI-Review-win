"""⑥ export：docx 双版本导出（设计文档 §5.2⑥）。

- 清洁版（{原名}_审校修订1_.docx）：accepted → 替换为 suggestion；custom → 替换为
  custom_text；rejected / pending → 保留原文。
- 留痕版（{原名}_审校修订2_.docx）：accepted → ｛～原文:{original} AI:{suggestion}～｝；
  custom → ｛～原文:{original} 用户:{custom_text}～｝（标记语义移植旧版
  update_text_change_both / update_text_change_self_define）。
- 替换锚点：correction.original 是句子片段，在其所属句子文本内做一次替换
  （blocks.text 由句子经 "\\n" 连接而成，句级替换与"以句子为锚点在块内替换一次"
  等价，且避免同文重复句的误伤）；锚点找不到时降级：记 warning 事件并保留原文
  （M4 已知问题——LLM 未逐字摘录 original 的场景）。
- 段落结构还原：block → 段落；chapter 标题行用 Heading 样式；表格占位符句
  （[{表格不予审校_N}]）独立成段，生成后按占位符把 tables.docx 中第 N 个表格
  还原回原位（移植旧 replace_placeholders_with_tables，改为 deepcopy + addprevious
  原位插入，两个版本互不干扰）。
- 全文段落首行缩进 docx.first_line_indent 英寸（默认 0.5，移植旧
  add_tab_indent_to_paragraphs）。
- 前置状态：pending_manual / manual_done / done；没有任何 accepted/custom 决定时
  仍导出（等同原文），adopted=0。成功后状态 → done，产物路径与 adopted 数记
  documents.exports_json 与 job_events。
"""
from __future__ import annotations

import json
import re
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any, Callable

from docx import Document as DocxDocument
from docx.shared import Inches
from sqlmodel import Session, select

from app.core.db import engine
from app.core.user_settings import first_line_indent_inches, output_dir
from app.models import Block, Correction, Document, Sentence
from app.pipeline.common import project_dir, set_document_status
from app.pipeline.ingest import PLACEHOLDER_RE as _PLACEHOLDER_SEARCH_RE

Emit = Callable[[str, dict], None]

EXPORTABLE_STATUSES = ("pending_manual", "manual_done", "done")

_PLACEHOLDER_LINE_RE = re.compile(r"^\[\{表格不予审校_(\d+)\}\]$")

# 导出文件名后缀（沿用旧版命名）
CLEAN_SUFFIX = "_审校修订1_.docx"
MARKED_SUFFIX = "_审校修订2_.docx"


def _noop_emit(_event: str, _data: dict) -> None:
    pass


def _marked_text(original: str, replacement: str, source_label: str) -> str:
    return f"｛～原文:{original} {source_label}:{replacement}～｝"


def _apply_sentence_corrections(
    sentence: Sentence,
    corrections: list[Correction],
    version: str,
    warn: Callable[[str], None],
) -> str:
    """在句子文本内按 id 序逐条应用已采纳决定；锚点缺失 → warn + 保留原文。"""
    text = sentence.text
    for c in corrections:
        if c.decision == "accepted":
            replacement, label = c.suggestion, "AI"
        elif c.decision == "custom":
            replacement, label = (c.custom_text or ""), "用户"
        else:  # rejected / pending → 保留原文
            continue
        if not replacement:
            warn(f"correction#{c.id} 替换文本为空（{c.decision}），已保留原文")
            continue
        if c.original not in text:
            warn(
                f"correction#{c.id} 锚点未命中（句子 {sentence.id}）："
                f"original={c.original[:40]!r} 不在句中，已保留原文"
            )
            continue
        new = replacement if version == "clean" else _marked_text(c.original, replacement, label)
        text = text.replace(c.original, new, 1)
    return text


def _build_paragraph_specs(
    session: Session, doc_id: str, version: str, warn: Callable[[str], None]
) -> list[dict[str, Any]]:
    """按 blocks 重建段落序列。

    返回 [{kind: heading|para|placeholder, text, table_index?}]：
    - chapter 标题（块内首句且与 block.chapter 同文）→ Heading 段落（应用修订后文本）；
    - 表格占位符句 → 独立占位段落（记录表格序号）；
    - 其余句子按块拼接为正文段落（句间无分隔符，还原原文连续文本），
      遇占位符切段；is_reference 块同样输出（不审校，原文照搬）。
    """
    blocks = session.exec(
        select(Block).where(Block.document_id == doc_id).order_by(Block.idx)
    ).all()
    specs: list[dict[str, Any]] = []
    for block in blocks:
        sentences = session.exec(
            select(Sentence).where(Sentence.block_id == block.id).order_by(Sentence.idx)
        ).all()
        buffer: list[str] = []

        def flush_buffer() -> None:
            if buffer:
                specs.append({"kind": "para", "text": "".join(buffer)})
                buffer.clear()

        for sentence in sentences:
            placeholder = _PLACEHOLDER_LINE_RE.match(sentence.text)
            if placeholder:
                flush_buffer()
                specs.append(
                    {
                        "kind": "placeholder",
                        "text": sentence.text,
                        "table_index": int(placeholder.group(1)),
                    }
                )
                continue
            corrections = session.exec(
                select(Correction)
                .where(Correction.sentence_id == sentence.id)
                .order_by(Correction.id)
            ).all()
            revised = _apply_sentence_corrections(sentence, corrections, version, warn)
            if sentence.idx == 0 and block.chapter and sentence.text == block.chapter:
                flush_buffer()
                specs.append({"kind": "heading", "text": revised})
            else:
                buffer.append(revised)
        flush_buffer()
    return specs


def _render_docx(specs: list[dict[str, Any]], dest_path: Path) -> None:
    """段落规格 → docx：Heading 样式标题 + 正文段落 + 占位符段落。"""
    doc = DocxDocument()
    for spec in specs:
        if spec["kind"] == "heading":
            doc.add_heading(spec["text"], level=1)
        else:
            doc.add_paragraph(spec["text"])
    doc.save(str(dest_path))


def _restore_tables(docx_path: Path, tables_path: Path) -> int:
    """把 [{表格不予审校_N}] 占位符段落替换回 tables.docx 中第 N 个表格。

    移植旧 replace_placeholders_with_tables；差异：deepcopy 表格节点 + addprevious
    原位插入（旧版按段落序号 insert，在已含表格的 body 中可能错位）。返回还原的表格数。
    """
    if not tables_path.exists():
        return 0
    doc = DocxDocument(str(docx_path))
    tables_doc = DocxDocument(str(tables_path))
    tables = tables_doc.tables
    restored = 0
    for para in list(doc.paragraphs):
        match = _PLACEHOLDER_SEARCH_RE.search(para.text)
        if match is None:
            continue
        index = int(match.group(1)) - 1
        if not (0 <= index < len(tables)):
            continue
        new_tbl = deepcopy(tables[index]._tbl)
        para._element.addprevious(new_tbl)
        para._element.getparent().remove(para._element)
        restored += 1
    if restored:
        doc.save(str(docx_path))
    return restored


def _apply_first_line_indent(docx_path: Path, inches: float) -> None:
    """全文段落首行缩进（移植旧 add_tab_indent_to_paragraphs，缩进量可配置）。"""
    if inches <= 0:
        return
    doc = DocxDocument(str(docx_path))
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.first_line_indent = Inches(inches)
    doc.save(str(docx_path))


def export_document(doc_id: str, emit: Emit | None = None) -> dict:
    """导出双版本 docx。成功：状态 → done 并记录 exports_json；失败：failed + 抛出。"""
    emit = emit or _noop_emit
    warnings: list[str] = []

    def warn(message: str) -> None:
        # 两个版本各构建一次段落，同一锚点缺失会产生重复告警——列表去重（事件仍逐次推送）
        if message not in warnings:
            warnings.append(message)
        emit("warning", {"message": message})

    with Session(engine) as session:
        doc = session.get(Document, doc_id)
        if doc is None:
            raise KeyError(f"文档不存在: {doc_id}")
        if doc.status not in EXPORTABLE_STATUSES:
            raise ValueError(f"当前状态 {doc.status} 不能导出，请先完成人工审校")
        filename = doc.filename
        original_status = doc.status

    emit("start", {"stage": "export", "filename": filename})
    try:
        pdir = project_dir(doc_id)
        tables_path = pdir / "tables.docx"

        # adopted = accepted + custom 决定数（锚点缺失降级不影响决定计数，见 warnings）
        with Session(engine) as session:
            blocks = session.exec(select(Block).where(Block.document_id == doc_id)).all()
            block_ids = [b.id for b in blocks]
            sentence_ids = (
                session.exec(select(Sentence.id).where(Sentence.block_id.in_(block_ids))).all()
                if block_ids
                else []
            )
            adopted = (
                len(
                    session.exec(
                        select(Correction).where(
                            Correction.sentence_id.in_(sentence_ids),
                            Correction.decision.in_(("accepted", "custom")),
                        )
                    ).all()
                )
                if sentence_ids
                else 0
            )
            specs = {
                version: _build_paragraph_specs(session, doc_id, version, warn)
                for version in ("clean", "marked")
            }

        out_dir_setting = output_dir()
        out_dir = Path(out_dir_setting) if out_dir_setting else pdir / "exports"
        out_dir.mkdir(parents=True, exist_ok=True)
        stem = Path(filename).stem or "document"
        paths = {
            "clean": out_dir / f"{stem}{CLEAN_SUFFIX}",
            "marked": out_dir / f"{stem}{MARKED_SUFFIX}",
        }

        indent = first_line_indent_inches()
        restored: dict[str, int] = {}
        for version, path in paths.items():
            _render_docx(specs[version], path)
            restored[version] = _restore_tables(path, tables_path)
            _apply_first_line_indent(path, indent)
            emit(
                "progress",
                {
                    "version": version,
                    "path": str(path),
                    "tables_restored": restored[version],
                    "paragraphs": len(specs[version]),
                },
            )

        exported_at = datetime.utcnow().isoformat()
        exports_payload = {
            "clean": str(paths["clean"]),
            "marked": str(paths["marked"]),
            "adopted": adopted,
            "warnings": len(warnings),
            "exported_at": exported_at,
        }
        with Session(engine) as session:
            row = session.get(Document, doc_id)
            if row is not None:
                row.exports_json = json.dumps(exports_payload, ensure_ascii=False)
                session.add(row)
                session.commit()
        set_document_status(doc_id, "done")
        emit(
            "done",
            {
                "clean_path": str(paths["clean"]),
                "marked_path": str(paths["marked"]),
                "adopted": adopted,
                "warnings": len(warnings),
            },
        )
        return {
            "clean_path": str(paths["clean"]),
            "marked_path": str(paths["marked"]),
            "adopted": adopted,
            "warnings": warnings,
            "tables_restored": restored,
            "status": "done",
        }
    except Exception as exc:
        # 导出失败不破坏审校成果：状态保持原状，仅记录错误（与 ingest/segment 的 failed 语义不同）
        set_document_status(doc_id, original_status, str(exc))
        emit("error", {"message": str(exc)})
        raise


def list_exports(doc: Document) -> list[dict]:
    """读取 documents.exports_json，返回产物列表（含存在性与大小）。"""
    if not doc.exports_json:
        return []
    try:
        payload = json.loads(doc.exports_json)
    except json.JSONDecodeError:
        return []
    items = []
    for kind, key, suffix in ((1, "clean", CLEAN_SUFFIX), (2, "marked", MARKED_SUFFIX)):
        raw = payload.get(key)
        if not raw:
            continue
        path = Path(raw)
        items.append(
            {
                "kind": kind,
                "label": "清洁版" if kind == 1 else "留痕版",
                "name": path.name,
                "path": str(path),
                "exists": path.exists(),
                "size": path.stat().st_size if path.exists() else None,
            }
        )
    return items
