import os  # 导入os模块，用于操作文件系统
import pypandoc  # 导入pypandoc模块，用于文件格式转换
from docx import Document  # 导入Document类，用于操作Word文档
import shutil  # 导入shutil模块，用于文件和文件夹的高级操作，如删除文件夹
import logging  # 导入logging模块，用于记录日志

# 设置日志记录的配置，定义日志级别和日志格式
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def delete_existing_files_and_folders(folder, exclude_files):
    # 删除指定文件夹中的文件和文件夹，除了排除列表中的项
    try:
        for item_name in os.listdir(folder):  # 列出文件夹中所有项
            item_path = os.path.join(folder, item_name)  # 获取每个项的完整路径
            if item_name not in exclude_files:  # 如果项不在排除列表中
                if os.path.isfile(item_path):  # 如果是文件，则删除
                    os.remove(item_path)
                    logging.info(f"已删除文件: {item_path}")  # 记录日志
                elif os.path.isdir(item_path):  # 如果是文件夹，则删除
                    shutil.rmtree(item_path)
                    logging.info(f"已删除文件夹: {item_path}")  # 记录日志
    except Exception as e:
        logging.error(f"删除文件和文件夹时出错: {e}")  # 记录错误日志

def replace_backslashes_in_md(md_file_path):
    # 在Markdown文件中替换所有反斜杠为正斜杠
    with open(md_file_path, 'r', encoding='utf-8') as file:  # 打开文件以读取
        content = file.read()  # 读取文件内容
    updated_content = content.replace('\\', '/')  # 替换反斜杠为正斜杠
    with open(md_file_path, 'w', encoding='utf-8') as file:  # 打开文件以写入
        file.write(updated_content)  # 写入修改后的内容

def add_indents_to_md(md_file_path, docx_file_path):
    # 在Markdown文件中添加缩进占位符
    with open(md_file_path, 'r', encoding='utf-8') as file:  # 打开Markdown文件以读取
        md_content = file.readlines()  # 读取所有行

    docx = Document(docx_file_path)  # 打开Word文档
    indents = []  # 初始化缩进信息列表
    for para in docx.paragraphs:  # 遍历Word文档中的段落
        # 获取段落的左缩进、首行缩进和悬挂缩进信息
        left_indent = para.paragraph_format.left_indent
        first_line_indent = para.paragraph_format.first_line_indent
        hanging_indent = -first_line_indent if first_line_indent and first_line_indent < 0 else 0
        indents.append((left_indent, first_line_indent, hanging_indent))  # 添加缩进信息到列表

    new_md_content = []  # 初始化修改后的Markdown内容列表
    indent_idx = 0  # 初始化索引计数器
    for line in md_content:  # 遍历Markdown文件的每一行
        if line.strip() != "" and not line.startswith("[^"):  # 如果行不为空且不是注释行
            if indent_idx < len(indents):  # 如果索引在缩进信息范围内
                left_indent, first_line_indent, hanging_indent = indents[indent_idx]  # 获取缩进信息
                indent_text = ''  # 初始化缩进文本
                # 根据缩进信息构建缩进占位符文本
                if left_indent:
                    indent_text += '[left_indent]'
                if first_line_indent:
                    indent_text += '[first_line_indent]'
                if hanging_indent:
                    indent_text += '[hanging_indent]'
                new_md_content.append(indent_text + line)  # 添加缩进占位符和行内容到新内容列表
                indent_idx += 1  # 索引加一
            else:
                logging.warning(f"Markdown中的段落比DOCX中的多。跳过行的缩进: {line.strip()}")  # 记录警告日志
                new_md_content.append(line)  # 添加行内容到新内容列表
        else:
            new_md_content.append(line)  # 直接添加行内容到新内容列表

    with open(md_file_path, 'w', encoding='utf-8') as file:  # 打开Markdown文件以写入
        file.writelines(new_md_content)  # 写入所有修改后的内容

def convert_file_md(source_file, output_file):
    # 将文件转换为Markdown格式，并处理相关缩进
    output_format = 'markdown'  # 设置输出格式为Markdown
    extract_media_path = os.path.splitext(output_file)[0] + '_media'  # 设置媒体文件的提取路径
    extra_args = ['--extract-media=' + extract_media_path, '--toc', '--standalone']  # 设置额外的pandoc参数
    output = pypandoc.convert_file(source_file, to=output_format, outputfile=output_file, extra_args=extra_args)  # 执行转换
    replace_backslashes_in_md(output_file)  # 替换反斜杠
    add_indents_to_md(output_file, source_file)  # 添加缩进占位符
    logging.info("MD文件格式转换成功")  # 记录日志

def replace_indent_placeholders(docx_file_path):
    # 替换Word文档中的缩进占位符为实际的制表符
    doc = Document(docx_file_path)  # 打开Word文档
    for para in doc.paragraphs:  # 遍历所有段落
        text = para.text  # 获取段落文本
        # 替换所有缩进占位符为制表符
        text = text.replace('[left_indent]', '\t').replace('[first_line_indent]', '\t').replace('[hanging_indent]', '\t')
        if text != para.text:  # 如果文本有变化
            para.clear()  # 清除原段落内容
            para.add_run(text)  # 添加新的段落内容
    doc.save(docx_file_path)  # 保存修改后的文档
    logging.info(f"Docx文件缩进占位符替换成功: {docx_file_path}")  # 记录日志

def convert_md_to_docx(input_md_path, output_docx_path):
    # 将Markdown文件转换为Word文档，并处理缩进占位符
    input_md_path = os.path.normpath(input_md_path)  # 规范化输入路径
    output_docx_path = os.path.normpath(output_docx_path)  # 规范化输出路径
    output = pypandoc.convert_file(input_md_path, 'docx', outputfile=output_docx_path)  # 执行转换
    replace_indent_placeholders(output_docx_path)  # 替换缩进占位符
    logging.info("Word文档格式转换成功")  # 记录日志
    return output

if __name__ == '__main__':
    test_folder = 'Doc_conver_test'  # 定义测试文件夹
    test_files = os.listdir(test_folder)  # 列出测试文件夹中的所有文件

    for file_name in test_files:  # 遍历测试文件夹中的文件
        source_file = os.path.join(test_folder, file_name)  # 获取源文件的完整路径
        if file_name.endswith('.docx') and not file_name.endswith('_converted.docx'):  # 如果文件是DOCX格式并且没有转换过
            exclude_files = [file_name]  # 设置排除列表
            delete_existing_files_and_folders(test_folder, exclude_files)  # 删除文件夹中的其他文件

            output_md_file = os.path.splitext(source_file)[0] + '.md'  # 设置输出Markdown文件路径
            output_docx_file = os.path.splitext(source_file)[0] + '_converted.docx'  # 设置输出DOCX文件路径

            convert_file_md(source_file, output_md_file)  # 转换为Markdown
            convert_md_to_docx(output_md_file, output_docx_file)  # 将Markdown转换为DOCX
