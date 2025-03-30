from docx import Document
from lxml import etree
import os
import glob
import re
from config import path_manager


# 移除第一个表格
def remove_first_table(doc_path):
    doc = Document(doc_path)

    if doc.tables:
        first_table = doc.tables[0]

        tbl_index = None
        for i, block in enumerate(doc.element.body):
            if block is first_table._element:
                tbl_index = i
                break

        if tbl_index is not None:
            doc.element.body.remove(doc.element.body[tbl_index])

    doc.save(doc_path)  # 无论是否有表格，都保存文档

    return f"已从 {doc_path} 中移除审查意见表。"

# 提取所有表格
def extract_tables_from_word(source_file, destination_file):
    source_doc = Document(source_file)
    destination_doc = Document()

    for table in source_doc.tables:
        tbl_xml = etree.tostring(table._tbl)
        new_tbl = etree.fromstring(tbl_xml)
        destination_doc._element.get_or_add_body().append(new_tbl)
        destination_doc.add_paragraph()

    destination_doc.save(destination_file)  # 无论是否有表格，都保存文档

    return f"所有表格已从 {source_file} 提取并保存到 {destination_file}。"

# 移除所有表格
def replace_tables(docx_path, modified_path):
    doc = Document(docx_path)
    table_counter = 1

    elements = list(doc.element.body)
    for element in elements:
        if element.tag.endswith('tbl'):
            index = doc.element.body.index(element)
            doc.element.body.remove(element)

            placeholder_text = f'[{{表格不予审校_{table_counter}}}]\n'
            placeholder_paragraph = doc.add_paragraph()
            placeholder_paragraph.add_run(placeholder_text)
            doc.element.body.insert(index, placeholder_paragraph._element)

            table_counter += 1

    doc.save(modified_path)  # 无论是否有表格，都保存文档

    return f"所有表格已从 {docx_path} 移除并保存到 {modified_path}。"

# 替换占位符为表格
def replace_placeholders_with_tables(doc_a_path, doc_b_path, doc_c_path):
    doc_a = Document(doc_a_path)
    doc_b = Document(doc_b_path)

    tables_b = doc_b.tables

    for i, para in enumerate(doc_a.paragraphs):
        # 使用转义字符对特殊字符进行转义
        match = re.search(r'\[\{表格不予审校_(\d+)\}\]', para.text)
        if match:
            placeholder_index = int(match.group(1)) - 1
            if placeholder_index < len(tables_b):
                p_element = para._element
                p_element.getparent().remove(p_element)

                table_to_insert = tables_b[placeholder_index]
                doc_a._element.body.insert(i, table_to_insert._tbl)

    doc_a.save(doc_c_path)  # 无论是否有表格，都保存文档











# 测试
if __name__ == "__main__":
    operate=1

    # 生成表处理文件路径
    def generate_file_paths(file_name):
        """生成文件处理所需的路径"""
        # 使用路径管理器生成文件路径
        paths = path_manager.generate_file_paths(file_name)
        return paths['begin_path'], paths['no_table'], paths['path_extract']

    path_1 = r"C:\Users\severin\OneDrive\CodeLib\ai_review\sug_0\_1_原文件\*"
    files = glob.glob(path_1)

    for file in files:
        path_extract, path_remove, path_replace = generate_file_paths(file)

        if operate==1:
            # 添加
            q=remove_first_table(file)
            a=extract_tables_from_word(file, path_extract)
            b=replace_tables(file, path_remove)
            c=replace_placeholders_with_tables(path_remove, path_extract, path_replace)
            print(q)
        elif operate==2:
            # 删除
            os.remove(path_extract)
            os.remove(path_remove)
            os.remove(path_replace)
        else:
            print('operate error')